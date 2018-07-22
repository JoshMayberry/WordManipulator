__version__ = "2.5.0"
##Does not use win32com

#Import standard elements
import os
import sys
import warnings

import docx

#Controllers
def build(*args, **kwargs):
	"""Starts the GUI making process."""

	return Excel(*args, **kwargs)

#Iterators
class Iterator(object):
	"""Used by handle objects to iterate over their nested objects."""

	def __init__(self, data, filterNone = False):
		if (not isinstance(data, (list, dict))):
			data = data[:]

		self.data = data
		if (isinstance(self.data, dict)):
			self.order = list(self.data.keys())

			if (filterNone):
				self.order = [key for key in self.data.keys() if key != None]
			else:
				self.order = [key if key != None else "" for key in self.data.keys()]

			self.order.sort()

			self.order = [key if key != "" else None for key in self.order]

	def __iter__(self):
		return self

	def __next__(self):
		if (not isinstance(self.data, dict)):
			if not self.data:
				raise StopIteration

			return self.data.pop()
		else:
			if not self.order:
				raise StopIteration

			key = self.order.pop()
			return self.data[key]

#Global Inheritance Classes
class Utilities():
	def __init__(self):
		"""Functions to make the Excel module easier.

		Example Input: Utilities()
		"""

		#Internal Variables
		self.childCatalogue = {} #{label (str): handle (child)}
		
	def __repr__(self):
		representation = f"{type(self).__name__}(id = {id(self)})"
		return representation

	def __str__(self):
		output = f"{type(self).__name__}()\n-- id: {id(self)}\n"
		if (hasattr(self, "parent") and (self.parent != None)):
			output += f"-- Parent: {self.parent.__repr__()}\n"
		return output

	def __len__(self):
		return len(self[:])

	def __contains__(self, key):
		return self._get(self.childCatalogue, key, returnExists = True)

	def __iter__(self):
		return Iterator(self.childCatalogue)

	def __getitem__(self, key):
		return self._get(self.childCatalogue, key)

	def __setitem__(self, key, value):
		self.childCatalogue[key] = value

	def __delitem__(self, key):
		del self.childCatalogue[key]

	def __enter__(self):
		return self

	def __exit__(self, exc_type, exc_value, traceback):
		self.entered = False
		if (traceback != None):
			print(exc_type, exc_value)
			return False

	def _get(self, itemCatalogue, itemLabel = None, returnExists = False):
		"""Searches the label catalogue for the requested object.

		itemLabel (any) - What the object is labled as in the catalogue
			- If slice: objects will be returned from between the given spots 
			- If None: Will return all that would be in an unbound slice

		Example Input: _get(self.childCatalogue)
		Example Input: _get(self.childCatalogue, 0)
		Example Input: _get(self.childCatalogue, slice(None, None, None))
		Example Input: _get(self.childCatalogue, slice(2, 7, None))
		"""

		#Account for retrieving all nested
		if (itemLabel == None):
			itemLabel = slice(None, None, None)

		#Account for indexing
		if (isinstance(itemLabel, slice)):
			if (itemLabel.step != None):
				raise FutureWarning(f"Add slice steps to _get() for indexing {self.__repr__()}")
			
			elif ((itemLabel.start != None) and (itemLabel.start not in itemCatalogue)):
				errorMessage = f"There is no item labled {itemLabel.start} in the row catalogue for {self.__repr__()}"
				raise KeyError(errorMessage)
			
			elif ((itemLabel.stop != None) and (itemLabel.stop not in itemCatalogue)):
				errorMessage = f"There is no item labled {itemLabel.stop} in the row catalogue for {self.__repr__()}"
				raise KeyError(errorMessage)

			handleList = []
			begin = False
			for item in sorted(itemCatalogue.keys()):
				#Allow for slicing with non-integers
				if ((not begin) and ((itemLabel.start == None) or (itemCatalogue[item].label == itemLabel.start))):
					begin = True
				elif ((itemLabel.stop != None) and (itemCatalogue[item].label == itemLabel.stop)):
					break

				#Slice catalogue via creation date
				if (begin):
					handleList.append(itemCatalogue[item])
			return handleList

		elif (itemLabel not in itemCatalogue):
			answer = None
		else:
			answer = itemCatalogue[itemLabel]

		if (returnExists):
			return answer != None

		if (answer != None):
			if (isinstance(answer, (list, tuple, range))):
				if (len(answer) == 1):
					answer = answer[0]
			return answer

		errorMessage = f"There is no item labled {itemLabel} in the data catalogue for {self.__repr__()}"
		raise KeyError(errorMessage)

	def getUnique(self, base = "{}", increment = 1, start = 1, exclude = []):
		"""Returns a unique name with the given criteria.

		Example Input: getUnique()
		Example Input: getUnique("Format_{}")
		Example Input: getUnique(exclude = [item.database_id for item in self.parent])
		"""

		if (not isinstance(exclude, (list, tuple, range))):
			exclude = [exclude]

		while True:
			ending = start + increment - 1
			if ((base.format(ending) in self) or (base.format(ending) in exclude) or (ending in exclude) or (str(ending) in [str(item) for item in exclude])):
				increment += 1
			else:
				break
		return base.format(ending)

class Utilities_Widget(Utilities):
	pass
	# def __exit__(self, *args, **kwargs):
	# 	answer = super().__exit__(*args, **kwargs)
	# 	self.apply()
	# 	return answer

#Handles
class Word(Utilities):
	def __init__(self):
		"""Works with word files.
		Documentation for docx can be found at: http://python-docx.readthedocs.io/en/latest/
		Example Input: Word()
		"""

		super().__init__()

	def new(self, label, *args, **kwargs):
		"""Creates a new document ans saves it in memmory.

		label (str) - The label of the workbook
		firstSheet (str) - The label for the first sheet in the workbook
			- If None: The workbook will start off without any sheets

		Example Input: new("test")
		"""

		document = self.Document(self, label, *args, **kwargs)
		self[label] = document

		return document

	def save(self, label, *args, **kwargs):
		"""Saves the document to a specified location.

		Example Input: save("test")
		"""

		self[label].save(*args, **kwargs)

	def load(self, label, *args, **kwargs):
		"""Loads a document from a specified location into memmory.

		Example Input: load("test")
		"""

		self[label].load(*args, **kwargs)

	def run(self, label, *args, **kwargs):
		"""Opens the ms word file for the user.

		Example Input: run("converted")
		"""

		self[label].run(*args, **kwargs)

	class Document(Utilities):
		def __init__(self, parent, label):
			"""A handle for the workbook.

			firstSheet (str) - The label for the first sheet in the workbook
				- If None: The workbook will start off without any sheets

			Example Input: Document(self, label)
			"""
			
			super().__init__()

			self.parent = parent
				
			if (label == None):
				label = self.getUnique("Document_{}")
			self.label = label
			self.title = None
			
			self.imageCatalogue = {} #(dict) - Used to catalogue all of the images in the document. {sheet title: [top-left corner cell (row, column), image as a PIL image]}

			self.load()
			self.setTitle()

			# if (firstSheet != None):
			# 	sheet = self.Sheet(self, firstSheet)
			# 	self[label] = sheet
			# 	self.select()

		def setTitle(self, title = None):
			"""Changes the title of the workbook.

			title (str) - The title of the workbook
				- If None: Will use the label for the workbook

			Example Input: setTitle("test")
			"""

			self.thing.core_properties.title = title or self.label

		def getTitle(self):
			"""Returns the title of the workbook.

			Example Input: getTitle()
			"""

			return self.thing.core_properties.title

		def setSubject(self, text = None):
			"""The topic of the content of the resource"""

			self.thing.core_properties.subject = text or ""

		def getVersion(self):
			return self.thing.core_properties.version

		def setVersion(self, text = None):
			"""The topic of the content of the resource"""

			self.thing.core_properties.version = text or ""

		def getSubject(self):
			return self.thing.core_properties.subject

		def setAuthor(self, text = None):
			"""An entity primarily responsible for making the content of the resource"""

			self.thing.core_properties.author = text or ""

		def getAuthor(self):
			return self.thing.core_properties.author

		def setCategory(self, text = None):
			"""A categorization of the content of this package. 
			Example values might include: Resume, Letter, Financial Forecast, Proposal, or Technical Presentation.
			"""

			self.thing.core_properties.category = text or ""

		def getCategory(self):
			return self.thing.core_properties.category

		def setComments(self, text = None):
			"""An account of the content of the resource"""

			self.thing.core_properties.comments = text or ""

		def getComments(self):
			return self.thing.core_properties.comments

		def setContentStatus(self, time = None):
			"""Completion status of the document, e.g. 'draft'"""

			self.thing.core_properties.content_status = text or ""

		def getContentStatus(self):
			return self.thing.core_properties.content_status

		def setIdentifier(self, time = None):
			"""An unambiguous reference to the resource within a given context, e.g. ISBN"""

			self.thing.core_properties.identifier = time or dateTime.now()

		def getIdentifier(self):
			return self.thing.core_properties.identifier

		def setKeywords(self, time = None):
			"""Descriptive words or short phrases likely to be used as search terms for this document"""

			self.thing.core_properties.keywords = text or ""

		def getKeywords(self):
			return self.thing.core_properties.keywords

		def setLanguage(self, time = None):
			"""Language the document is written in"""

			self.thing.core_properties.language = text or ""

		def getLanguage(self):
			return self.thing.core_properties.language

		def setRevision(self, value = None):
			"""Number of this revision, incremented by Word each time the document is saved. 
			Note however python-docx does not automatically increment the revision number when it saves a document.
			"""

			self.thing.core_properties.revision = value or self.getRevision() + 1

		def getRevision(self):
			return self.thing.core_properties.revision

		def setTime_created(self, time = None):
			"""Time of intial creation of the document"""

			self.thing.core_properties.created = time or dateTime.now()

		def getTime_created(self):
			return self.thing.core_properties.created

		def setTime_printed(self, time = None):
			"""Time the document was last printed"""

			self.thing.core_properties.last_printed = time or dateTime.now()

		def getTime_printed(self):
			return self.thing.core_properties.last_printed

		def setTime_modified(self, time = None):
			"""Time the document was last modified"""

			self.thing.core_properties.modified = time or dateTime.now()

		def getTime_modified(self):
			return self.thing.core_properties.modified

		def setLastModifiedBy(self, time = None):
			"""Name or other identifier (such as email address) of person who last modified the document"""

			self.thing.core_properties.last_modified_by = text or ""

		def getLastModifiedBy(self):
			return self.thing.core_properties.last_modified_by

		def save(self, filePath = "", overlayOk = True, temporary = False, saveImages = True):
			"""Saves the workbook to a specified location.

			filePath (str)   - Where the file is located
			overlayOk (bool) - If True: Images can overlap. If False: Any images under otehr ones will be deleted. If None: Images will be scooted to the right until they are ont under another
			temporary (bool) - If True: The file will be saved under the same name, but with "_temp" after it. For debugging things
			saveImages (bool) - If True: Images in the document will be preserved upon loading
				Images, charts, etc. are not read by openpyxl.
				In order to preserve images, charts, etc., each image is loaded and re-written into the loaded workbook
				Method for preservation from http://www.penwatch.net/cms/?p=582
				Help from: code.activestate.com/recipes/528870-class-for-writing-content-to-excel-and-formatting

			Example Input: save()
			"""

			if (temporary):
				fileName += "_temp"
			else:
				fileName = self.label

			try:
				#Ensure correct format
				if ("." not in fileName):
					fileName += ".docx"

				self.thing.save(os.path.join(filePath, fileName))
			
			except IOError:
				#A book by that name is already open
				print("ERROR: The word file is still open. The file has still been saved. Just close the current file without saving.")

		def load(self, filePath = None, readImages = False):
			"""Loads a workbook from a specified location into memmory.

			filePath (str) - Where the file is located
				- If None: Will create a new, blank document
			readImages (bool) - If True: Images in the document will be preserved upon loading
				Images, charts, etc. are not read by openpyxl.
				In order to preserve images, charts, etc., each image is loaded and re-written into the loaded workbook
				Method for preservation from http://www.penwatch.net/cms/?p=582
				Help from: code.activestate.com/recipes/528870-class-for-writing-content-to-excel-and-formatting

			Example Input: load()
			"""

			fileName = self.label

			#Ensure correct format
			if ("." not in fileName):
				fileName += ".xlsx"

			#Load the workbook into memory
			self.thing = docx.Document(docx = filePath)
			self.update()

		def run(self, filePath = "./"):
			"""Opens the word file for the user.

			filePath (str) - Where the file is located

			Example Input: run()
			"""

			#Ensure correct format
			if ("." not in fileName):
				fileName += ".xlsx"

			try:
				os.startfile(os.path.join(filePath, fileName))
			except AttributeError:
				subprocess.call(['open', fileName])

		def update(self):
			self.updateSections()

		def updateSections(self):
			for i, sectionObject in enumerate(self.thing.sections):
				if (i not in self):
					self.addSection(thing = sectionObject)
				elif (sectionObject is not self[i].thing):
					self[i].thing = sectionObject

		def getSection(self, index = None):
			if (index != None):
				if (index > 0):
					return self[index]
				else:
					sectionList = self[:]
					return sectionList[index]
			return self[:]

		def addParagraph(self, *args, **kwargs):
			"""Adds a paragraph to the document."""

			return self.Paragraph(self, *args, **kwargs)

		def addHeader(self, level = 0, *args, **kwargs):
			"""Alias for a header paragraph"""

			return self.addParagraph(header = level)

		def addIntense(self, *args, **kwargs):
			"""Alias for an intense quote paragraph"""

			return self.addParagraph(intense = True)

		def addList(self, bullet = True, *args, **kwargs):
			"""Alias for a list paragraph"""

			return self.addParagraph(bulletList = bullet)

		def addImage(self, *args, **kwargs):
			"""Adds an image to the document."""

			return self.Image(self, *args, **kwargs)

		def addSection(self, *args, **kwargs):
			"""Adds a section to the document."""

			return self.Section(self, *args, **kwargs)

		def addTable(self, *args, **kwargs):
			"""Adds a table to the document."""

			return self.Table(self, *args, **kwargs)

		def addPageBreak(self):
			"""Adds a page break to the document."""

			self.thing.add_page_break()

		class Paragraph(Utilities_Widget):
			def __init__(self, parent, header = None, intense = None, bulletList = None, thing = None):

				#Intitialize Inherited Modules
				super().__init__()

				#Internal Variables
				self.parent = parent

				if (thing != None):
					self.thing = thing
				else:
					if (header != None):
						self.thing = self.parent.thing.add_heading(level = header)
					elif (intense != None):
						self.thing = self.parent.thing.add_paragraph(style = "Intense Quote")
					elif (bulletList != None):
						if (bulletList):
							self.thing = self.parent.thing.add_paragraph(style = "List Bullet")
						else:
							self.thing = self.parent.thing.add_paragraph(style = "List Number")
					else:
						self.thing = self.parent.thing.add_paragraph()

			def addText(self, text = "", bold = None, italic = None, underline = None):
				"""Adds text to the paragraph.

				Example Input: addText("Lorem Ipsum")
				Example Input: addText("Lorem Ipsum", bold = True, italic = True)
				Example Input: addText("Lorem Ipsum", bold = True, italic = True)
				"""

				segment = self.thing.add_run(text)

				if (bold != None):
					segment.bold = bold
				if (italic != None):
					segment.italic = italic
				if (underline != None):
					segment.underline = underline

		class Image(Utilities_Widget):
			def __init__(self, parent, filePath, width = None, height = None, thing = None):

				#Intitialize Inherited Modules
				super().__init__()

				#Internal Variables
				self.parent = parent

				if (width != None):
					width = docx.shared.Inches(width)
				if (height != None):
					width = docx.shared.Inches(height)

				if (thing != None):
					self.thing = thing
				else:
					self.thing = self.parent.thing.add_picture(filePath, width = width, height = height)

		class Section(Utilities_Widget):
			def __init__(self, parent, thing = None):

				#Intitialize Inherited Modules
				super().__init__()

				#Internal Variables
				self.parent = parent

				if (thing != None):
					self.thing = thing
				else:
					self.thing = self.parent.thing.add_section()

				#Nest section in document
				self.parent[len(self.parent)] = self

			def startOdd(self):
				"""Section begins on next odd page"""
				
				self.thing.start_type = docx.enum.section.WD_SECTION.ODD_PAGE

			def startEven(self):
				"""Section begins on next even page"""
				
				self.thing.start_type = docx.enum.section.WD_SECTION.EVEN_PAGE

			def startNew(self):
				"""Section begins on next new page"""
				
				self.thing.start_type = docx.enum.section.WD_SECTION.NEW_PAGE

			def startNewColumn(self):
				"""Section begins on next new column"""
				
				self.thing.start_type = docx.enum.section.WD_SECTION.NEW_COLUMN

			def startNone(self):
				"""Section begins after the last section"""
				
				self.thing.start_type = docx.enum.section.WD_SECTION.CONTINUOUS

			def setSize(self, width = None, height = None):
				self.setWidth(width)
				self.setHeight(height)

			def setWidth(self, value = None):
				"""Total page width used for this section, inclusive of all edge spacing values such as margins. 
				Page orientation is taken into account, so for example, its expected value would be Inches(11) for letter-sized paper when orientation is landscape.
				"""

				if (value != None):
					self.thing.page_width = docx.shared.Inches(value)

			def setHeight(self, value = None):
				"""Total page height used for this section, inclusive of all edge spacing values such as margins. 
				Page orientation is taken into account, so for example, its expected value would be Inches(8.5) for letter-sized paper when orientation is landscape.
				"""

				if (value != None):
					self.thing.page_height = docx.shared.Inches(value)

			def setHeight_header(self, value = None):
				"""Length object representing the distance from the top edge of the page to the top edge of the header. 
				None if no setting is present in the XML.
				"""

				if (value != None):
					self.thing.header_distance = docx.shared.Inches(value)

			def setHeight_footer(self, value = None):
				"""Length object representing the distance from the bottom edge of the page to the bottom edge of the footer. 
				None if no setting is present in the XML.
				"""

				if (value != None):
					self.thing.footer_distance = docx.shared.Inches(value)

			def vertical(self, state = True):
				if (state):
					self.thing.orientation = docx.enum.section.WD_ORIENT.PORTRAIT
				else:
					self.horizontal()

			def horizontal(self, state = True):
				if (state):
					self.thing.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
				else:
					self.vertical()

			def setMargins(self, left = None, right = None, top = None, bottom = None):
				self.setMargins_left(left)
				self.setMargins_right(right)
				self.setMargins_top(top)
				self.setMargins_bottom(bottom)

			def setMargins_left(self, value = None):
				"""Length object representing the left margin for all pages in this section in English Metric Units."""

				if (value != None):
					self.thing.left_margin = docx.shared.Inches(value)

			def setMargins_right(self, value = None):
				"""Length object representing the right margin for all pages in this section in English Metric Units."""

				if (value != None):
					self.thing.right_margin = docx.shared.Inches(value)

			def setMargins_top(self, value = None):
				"""Length object representing the top margin for all pages in this section in English Metric Units."""

				if (value != None):
					self.thing.top_margin = docx.shared.Inches(value)

			def setMargins_bottom(self, value = None):
				"""Length object representing the bottom margin for all pages in this section in English Metric Units."""
				
				if (value != None):
					self.thing.bottom_margin = docx.shared.Inches(value)

			def setGutter(self, value = None):
				"""Length object representing the page gutter size in English Metric Units for all pages in this section.
				The page gutter is extra spacing added to the inner margin to ensure even margins after page binding.
				"""
				
				if (value != None):
					self.thing.gutter = docx.shared.Inches(value)

		class Table(Utilities_Widget):
			def __init__(self, parent, rows = 1, columns = 1, thing = None):

				#Intitialize Inherited Modules
				super().__init__()

				#Internal Variables
				self.parent = parent

				if (thing != None):
					self.thing = thing
				else:
					self.thing = selfparent.thing.add_table(rows, columns)

			def getCells(self):
				pass
		
if (__name__ == "__main__"):
	word = Word()

	with word.new("test_2") as myDocument:
		with myDocument.getSection(-1) as mySection:
			mySection.setMargins(0.5, 0.5, 0.5, 0.5)

			with myDocument.addHeader() as myHeader:
				myHeader.addText("Document Title")

			with myDocument.addParagraph() as myParagraph:
				myParagraph.addText("A plain paragraph having some ")
				myParagraph.addText("bold text", bold = True)
				myParagraph.addText(" and some ")
				myParagraph.addText("italic text.", italic = True)

			with myDocument.addHeader(level = 1) as myHeader:
				myHeader.addText("Sub Heading")

			with myDocument.addIntense() as myParagraph:
				myParagraph.addText("Intense Quote")

			# with myDocument.addList(bullet = True) as myList:
			# 	myList.addItem("First item in unordered list")
			# 	myList.addItem("Second item in unordered list")

			# with myDocument.addList(bullet = False) as myList:
			# 	myList.addItem("First item in ordered list")
			# 	myList.addItem("Second item in ordered list")

			# myDocument.addImage("C:/Users/Kade/Pictures/Untitled.png", width = 1.25)
			myDocument.addPageBreak()

			# with myDocument.addTable(rows = 1, columns = 3) as myTable:
			# 	for i, cell in enumerate(myTable.getCells()):
			# 		cell.setText(f"Lorem {i}")

		myDocument.save()